from docx import Document
from docx.shared import RGBColor,Pt,Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import sys
import os
import random

document = Document()
# font = document.styles['Normal'].font

# font.name = 'RunWild-Demo'
#font.name = "Times New Roman"

txt = "1 2 3 4 This is the 13th article in my series of articles on Python for NLP. In the previous article, we saw how to create a simple rule-based chatbot that uses cosine similarity between the TF-IDF vectors of the words in the corpus and the user input, to generate a response. The TF-IDF model was basically used to convert word to numbers.\n In this article, we will study another very useful model that converts text to numbers i.e. the Bag of Words (BOW).\n\nSince most of the statistical algorithms, e.g machine learning and deep learning techniques, work with numeric data, therefore we have to convert text into numbers. Several approaches exist in this regard. However, the most famous ones are Bag of Words, TF-IDF, and word2vec. Though several libraries exist, such as Scikit-Learn and NLTK, which can implement these techniques in one line of code, it is important to understand the working principle behind these word embedding techniques. The best way to do so is to implement these techniques from scratch in Python and this is what we are going to do today."
t=document.add_paragraph()
for i in txt:
	rno = random.randint(0,2)
	k = t.add_run(i)
	k.bold = True
	font = k.font
	if rno==0:
		if i >='A' and i<='z' or i==' ':
			font.name = 'Run Wild - Demo'
		else:
			font.name = 'kristi'
	# elif rno==1:
		# font.name="kristi"
	else:
		font.name="LeviPen"

	font.color.rgb = RGBColor(0,15,85)
	# font.color.brightness = 0.1
document.save("file.docx")
